import json
import random
import re
from pprint import pprint

import yaml
import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Mm
from openpyxl import Workbook, load_workbook
from openpyxl.styles import PatternFill
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

AUDS = {
    "В-02": 0,
    "В-201": 0,
    "В-202": 0,
    "В-214": 0,
    "В-301": 0,
    "В-302": 0,
    "В-304": 0,
    "Г-102": 0,
    "Г-105": 0,
    "Г-110б": 0,
    "Г-206": 0,
    "Г-212": 0,
    "Г-219": 0,
    "Г-314": 0,
    "И-202": 0,
    "И-305": 0,
    "И-307": 0,
    "И-307a": 0,
    "И-309": 0,
    "К-207": 0,
    "Конференц-зал": 0,

    'K-301': 1,
    "K-302": 1,
    "K-303": 1,
    "K-305": 1,
    "K-311": 1,
    "K-313": 1,
    "K-315": 1,
    "K-315а": 1,
    "K-317": 1,
    "И-311": 1,
    "И-315": 1,
    "И-317": 1,
    "И-319": 1,
    "Г-амф.": 1,
    "Ж-216": 1,
    "Ж-319": 1,
    "Ж-321": 1,
    "Ж-323": 1,
    "Ж-амфитеатр": 1,
}


def process_file(document: Document, file_name):
    data = []
    table = document.tables[0]
    for row in table.rows[1:]:
        cells = list(row.cells)
        if cells[2].text.strip():
            dt = datetime.strptime(cells[2].text.strip() + " " + cells[3].text.strip(), "%d.%m.%Y %H.%M")

            data.append({
                'file_name': file_name,
                'title': cells[0].text,
                'type': cells[1].text,
                'date': dt,
                'aud': cells[4].text,
            })
        else:
            if cells[0].text and cells[2].text.strip():
                print(f"Остутсвует дата для {cells[0].text} в файле {document}")
    return data


def fill_data_file():
    """
    Формирует data.yaml из docx документов в папке ./data/2020
    :return:
    """
    data = {}
    for file in os.listdir("./data/2021"):
        if file.endswith(".docx"):
            title = file.replace('.docx', '')
            pth = os.path.join("./data/2021", file)
            data[title] = process_file(Document(pth), title)

    with open("data.yaml", "w", encoding="utf8") as f:
        yaml.dump(data, f, allow_unicode=True)


def set_cell_value(cell, value, highlight=False):
    if len(cell.paragraphs) == 0:
        cell.add_paragraph()
    par = cell.paragraphs[0]

    if len(par.runs) == 0:
        par.add_run()

    par.runs[0].font.size = Pt(12)
    if highlight:
        par.runs[0].font.highlight_color = highlight
    par.runs[0].text = value


def calculate_schedule():
    """
    Рассчет расписания по data.yaml
    :return:
    """
    with open("data.yaml", 'r', encoding="utf8") as f:
        data = yaml.load(f)

    random.seed(1)

    # словарик под расписание аудтиорий
    auds_schedule = {
        aud: {}
        for aud, v in AUDS.items()
    }

    # тусуем аудитории чтобы усреднить результат
    all_aud_ids = list(AUDS.keys())
    random.shuffle(all_aud_ids)

    # расстояние между занятиями
    gap = 5

    # тусуем занятия для проектирования
    all_items_to_schedule = []
    for key, items in data.items():
        all_items_to_schedule.extend(items)
    random.shuffle(all_items_to_schedule)
    all_items_to_schedule = sorted(all_items_to_schedule, key=lambda x: not(x['file_name'].startswith('Архитектуры, строительства')))

    # конференц-зал оставляем
    for item in all_items_to_schedule:
        if item['aud'] != 'Конференц-зал':
            item['aud'] = None
        else:
            dates = auds_schedule[item['aud']]
            date_items = dates.setdefault(item['date'].date(), [])
            date_items.append(item)

    # проектируем
    for item in all_items_to_schedule:
        found_aud = False

        if item['aud']:
            continue

        # расставляем занятия по приоритетами
        for priority in range(max(AUDS.values()) + 1):
            # формируем список аудиторий с указаным приоритетам
            aud_ids = [k for k in all_aud_ids if AUDS[k] == priority]
            if item['file_name'].startswith('Архитектуры, строительства'):
                aud_ids = sorted(aud_ids, key=lambda k: not k.startswith("Г"))

            # ищем свбодные аудитории
            for aud in aud_ids:
                dates = auds_schedule[aud]
                date_items = dates.setdefault(item['date'].date(), [])
                # если нашли, то фиксируем
                if not date_items:
                    date_items.append(item)
                    item['aud'] = aud
                    found_aud = True
                    break

            # если не нашли, пытаемся запихнуть занятия второй сменной
            if not found_aud:
                for aud in aud_ids:
                    dates = auds_schedule[aud]
                    date_items = dates.setdefault(item['date'].date(), [])

                    if all(abs((item['date'] - i['date']).total_seconds()) / 3600 >= gap for i in date_items):
                        date_items.append(item)
                        item['aud'] = aud
                        found_aud = True
                        break
            # если нашли занятия то нет смысла пытатся влезть в аудитории с более низким приоритетом
            if found_aud:
                break

        if not found_aud:
            print(f"Не удалось найти аудитории для {item['title']}")

    return data


def create_docx_documents(data):
    """
    Формирование docx документов по составленому расписанию
    :param data:
    :return:
    """
    for key, items in data.items():
        doc = Document("template.docx")
        table = doc.tables[0]

        previous_title = None
        previous_row_index = 1
        row_index = 0

        for row_index, item in enumerate(items, 1):
            row = table.add_row()

            cells = list(row.cells)

            title = item['title'].strip()
            if title != previous_title:
                if previous_title:
                    cell_start = table.rows[previous_row_index].cells[0]
                    cell_end = table.rows[row_index - 1].cells[0]
                    cell_start.merge(cell_end)
                    set_cell_value(cell_start, previous_title)

                    previous_row_index = row_index

                previous_title = title

            for index, value in enumerate([
                item['type'].strip(),
                "{:%d.%m.%Y}".format(item['date']),
                "{:%H.%M}".format(item['date']),
                item['aud'],
            ], 1):
                set_cell_value(cells[index], value, {
                    1: WD_COLOR_INDEX.BRIGHT_GREEN,
                    2: WD_COLOR_INDEX.YELLOW,
                    3: WD_COLOR_INDEX.RED,
                }.get(AUDS.get(item['aud'])))

            row_index += 1

        if previous_row_index != row_index:
            cell_start = table.rows[previous_row_index].cells[0]
            cell_end = table.rows[row_index - 1].cells[0]
            cell_start.merge(cell_end)
            set_cell_value(cell_start, previous_title)

        doc.save(os.path.join("output", f"{key}.docx"))


def generate_auds_schedule_document(data, use_real_auds=False):
    """
    Формирование сетки аудиторий в формате дата/аудитория
    :param data:
    :return:
    """
    result = {}

    for key, items in data.items():
        for item in items:
            date_item = result.setdefault(item['date'].date(), {})
            all_items = date_item.setdefault(item['aud'], [])
            all_items.append(item)

    wb = Workbook()
    ws = wb.active

    auds = []
    if use_real_auds:
        for dt, value in result.items():
            auds.extend(value.keys())
        auds = sorted(list(set(auds)))
    else:
        auds = sorted(AUDS.keys(), key=lambda x: x.replace("K", "К").replace("B", "В"))

    for date_index, date in enumerate(sorted(result.keys())):
        ws.cell(date_index + 2, 1, "{:%d.%m.%Y}".format(date))
        for aud_index, aud in enumerate(auds):
            ws.cell(1, aud_index + 2, f"{aud}")
            items = result.get(date, {}).get(aud, [])
            if items:
                title = "\n\n".join(
                    f"{'{:%H.%M}'.format(i['date'])}: {i['title']}" for i in sorted(items, key=lambda x: x['date']))
                cell = ws.cell(date_index + 2, aud_index + 2, f"{title}")
                cell.fill = PatternFill(fill_type='solid', start_color="FFDDDDDD", end_color="FFDDDDDD")

    wb.save("auds.xlsx")


def generate_auds_docx():
    def set_repeat_table_header(row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    wb = load_workbook("auds_15.06.2021.xlsx")
    ws = wb.active

    rows = list(ws.rows)
    first_row = rows[0]

    data = {}

    for cell in first_row[1:]:
        data[cell.value] = []

    for row in rows[1:]:
        for index, cell in enumerate(row):
            aud_name = first_row[index].value
            if index == 0:
                if isinstance(cell.value, datetime):
                    date = "{:%d.%m.%Y}".format(cell.value)
                else:
                    date = cell.value
            elif cell.value:
                data[aud_name].append({
                    'date': date,
                    'value': cell.value.strip().replace("\n\n\n", "#")
                        .replace("\n", " ")
                        .replace(r"/ ", " / ").replace("#", "\n"),
                })

    pprint(data)

    doc = Document("template2.docx")
    for aud_name, rows in data.items():
        if not rows:
            continue

        table = doc.add_table(0, 2, "Table Grid")
        table.columns[0].width = Mm(25)
        table.columns[1].width = Mm(150)

        row1 = table.add_row()
        row1.cells[0].merge(row1.cells[1])

        p = row1.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.text = aud_name
        r.font.bold = True
        r.font.size = Pt(16)

        set_repeat_table_header(row1)

        # row1 = table.rows[0]
        for r in rows:
            row = table.add_row()
            cells = list(row.cells)
            cells[0].text = r['date']
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # cells[0].width = Mm(22)
            cells[1].text = r['value']
            # cells[1].width = Mm(130)
        doc.add_page_break()

    doc.save("расписание_гос.экзаменов_по_аудиториям.docx")

def generate_script(data, use_real_auds=False, skip_not_exists_auds=False):
    """
     Формирование сетки аудиторий в формате дата/аудитория
     :param data:
     :return:
     """
    result = {}

    for key, items in data.items():
        for item in items:
            date_item = result.setdefault(item['date'].date(), {})
            all_items = date_item.setdefault(item['aud'], [])
            all_items.append(item)

    auds = []
    if use_real_auds:
        for dt, value in result.items():
            auds.extend(value.keys())
        auds = sorted(list(set(auds)))
    else:
        auds = sorted(AUDS.keys(), key=lambda x: x.replace("K", "К").replace("B", "В"))

    with open("auds_ids.json", encoding="utf8") as f:
        aud_info = json.load(f)
        aud_info = {i['obozn']: i for i in aud_info}

    not_found_auds = []
    for aud in auds:
        if aud not in aud_info:
            not_found_auds.append(aud)

    if not skip_not_exists_auds and not_found_auds:
        raise Exception(not_found_auds)

    sql_items = []
    for date_index, date in enumerate(sorted(result.keys())):
        for aud_index, aud in enumerate(auds):
            items = result.get(date, {}).get(aud, [])
            if items:
                i = items[0]
                if i['aud'] in aud_info:
                    title = i['title'].replace('\n', '\\n')
                    sql_items.append(f"(ARRAY[{aud_info[i['aud']]['id']}], 'ГЭК - {title}', '{i['date']:%Y-%m-%d}', '{{1,2,3,4,5,6}}', 100, '')")

    s = "INSERT INTO queries(auds, description, dt, pairs, type, in_charge) VALUES "
    s += ",\n".join(sql_items)
    with open("query.sql", 'w', encoding="utf8") as f:
        f.write(s)


def main():
    # generate_auds_docx()
    # fill_data_file()
    # data = calculate_schedule()
    # with open("processed.yaml", "w", encoding='utf8') as f:
    #     yaml.dump(data, f, allow_unicode=True)

    with open("data.yaml", encoding='utf8') as f:
        data = yaml.load(f)

    # create_docx_documents(data)
    # generate_auds_schedule_document(data, use_real_auds=True)
    generate_script(data, use_real_auds=True, skip_not_exists_auds=True)


if __name__ == '__main__':
    main()
